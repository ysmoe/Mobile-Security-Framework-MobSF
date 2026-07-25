# -*- coding: utf_8 -*-
"""
Shared Functions.

DOCX Report Generation
======================

Generate a Microsoft Word (.docx) report from the same Django template
context used by the PDF report, but using ``python-docx`` instead of
``wkhtmltopdf`` so the user's Word processor renders the file natively
and edits / translations stay editable in the open XML format.

Pipeline (parallels :mod:`mobsf.StaticAnalyzer.views.common.pdf`):

    DB -> handle_pdf_android/ios/win -> template.render(context)
                                                 |
                                  +--------------+--------------+
                                  |                             |
                                  v                             v
                          pdfkit.from_string              html_to_docx
                                  |                             |
                                  v                             v
                                PDF                          DOCX
"""
import io
import json
import logging
import os
import platform
import re
from html import unescape
from urllib.parse import urlparse
from urllib.request import url2pathname

from bs4 import BeautifulSoup
from bs4.element import NavigableString
from django.http import HttpResponse
from django.template.loader import get_template

import mobsf.MalwareAnalyzer.views.VirusTotal as VirusTotal
from mobsf.MobSF import settings
from mobsf.MobSF.utils import (
    is_md5,
    print_n_send_error_response,
)
from mobsf.StaticAnalyzer.models import (
    RecentScansDB,
    StaticAnalyzerAndroid,
    StaticAnalyzerIOS,
    StaticAnalyzerWindows,
)
from mobsf.StaticAnalyzer.views.common.appsec import (
    get_android_dashboard,
    get_ios_dashboard,
)
from mobsf.StaticAnalyzer.views.common.shared_func import (
    get_avg_cvss,
)
from mobsf.StaticAnalyzer.views.android.db_interaction import (
    get_context_from_db_entry as adb)
from mobsf.StaticAnalyzer.views.ios.db_interaction import (
    get_context_from_db_entry as idb)
from mobsf.StaticAnalyzer.views.windows.db_interaction import (
    get_context_from_db_entry as wdb)
# Register the relative_path / key / pathify template filters defined in
# static_analyzer. Importing this module is required because the filter
# registration attaches to ``django.template.defaulttags.register`` which
# is a process-global Library; the PDF template (``android_report.html``)
# uses ``{{ ... | relative_path }}`` so we must ensure the filter is in
# place when the same template is rendered for DOCX output.
from mobsf.StaticAnalyzer.views.android import static_analyzer as _android_sa  # noqa: F401,E501
from mobsf.StaticAnalyzer.views.ios import static_analyzer as _ios_sa  # noqa: F401
from mobsf.MobSF.views.authentication import (
    login_required,
)
from mobsf.StaticAnalyzer.views.common.pdf import (
    handle_pdf_android,
    handle_pdf_ios,
    handle_pdf_win,
)

logger = logging.getLogger(__name__)
ctype = 'application/json; charset=utf-8'

# Optional dependency. We import lazily so that the absence of python-docx
# in non-reporting paths doesn't break the rest of MobSF.
try:
    from docx import Document
    from docx.enum.table import WD_ALIGN_VERTICAL
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.shared import Inches, Pt, RGBColor
    DOCX_AVAILABLE = True
    IMPORT_ERROR = None
except ImportError as exc:  # pragma: no cover - environment guard
    DOCX_AVAILABLE = False
    IMPORT_ERROR = str(exc)
    logger.warning(
        'python-docx is not installed. DOCX report generation is disabled: %s',
        exc,
    )


# ---------------------------------------------------------------------------
# Inline resource helpers
# ---------------------------------------------------------------------------
def _local_path_from_url(url):
    """Translate a ``file://`` URL produced by the PDF pipeline into an
    absolute filesystem path that ``python-docx`` can read."""
    parsed = urlparse(url)
    if parsed.scheme != 'file':
        return None
    netloc = parsed.netloc or ''
    path = url2pathname(parsed.path)
    if netloc:
        # UNC path on Windows: file://server/share/x.png
        path = f'//{netloc}{path}'
    return os.path.normpath(path)


def _inline_image_bytes(src, base_url=None):
    """Resolve an ``<img src="...">`` value to raw bytes.

    Supports the three kinds of references the PDF templates emit:
      * ``data:image/png;base64,XXX``  (already-embedded diagrams)
      * ``file:///abs/path.png``      (PDF pipeline convention)
      * ``http(s)://...``             (unlikely, but we still try)
    """
    if not src:
        return None
    src = src.strip()
    if src.startswith('data:image'):
        m = re.match(r'data:image/[^;]+;base64,(.*)$', src, re.DOTALL)
        if m:
            import base64
            try:
                return base64.b64decode(m.group(1))
            except Exception:
                logger.warning('Invalid base64 image data')
                return None
        return None
    if src.startswith('file://'):
        path = _local_path_from_url(src)
        if path and os.path.isfile(path):
            try:
                with open(path, 'rb') as fh:
                    return fh.read()
            except OSError as exc:
                logger.warning('Cannot read file:// image %s: %s', path, exc)
        return None
    if base_url and not src.lower().startswith(('http://', 'https://', '/')):
        # Relative URL - try resolving against base_url
        candidate = os.path.join(base_url, src)
        if os.path.isfile(candidate):
            try:
                with open(candidate, 'rb') as fh:
                    return fh.read()
            except OSError:
                pass
    return None


# ---------------------------------------------------------------------------
# Style helpers
# ---------------------------------------------------------------------------
CHINESE_FONT = 'Noto Sans CJK SC'  # provided by fonts-noto-cjk in the image


def _set_run_font(run, size=11, bold=False, color=None):
    """Apply CJK-safe font + size to a run."""
    run.font.name = CHINESE_FONT
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn('w:rFonts'))
    if rfonts is None:
        from docx.oxml import OxmlElement
        rfonts = OxmlElement('w:rFonts')
        rpr.append(rfonts)
    rfonts.set(qn('w:eastAsia'), CHINESE_FONT)
    rfonts.set(qn('w:ascii'), CHINESE_FONT)
    rfonts.set(qn('w:hAnsi'), CHINESE_FONT)
    run.font.size = Pt(size)
    run.font.bold = bool(bold)
    if color is not None:
        run.font.color.rgb = RGBColor(*color)


def _bootstrap_document():
    """Create a Document with a sensible default style for CJK rendering."""
    doc = Document()
    style = doc.styles['Normal']
    style.font.name = CHINESE_FONT
    style.font.size = Pt(11)
    rpr = style.element.get_or_add_rPr()
    rfonts = rpr.find(qn('w:rFonts'))
    if rfonts is None:
        from docx.oxml import OxmlElement
        rfonts = OxmlElement('w:rFonts')
        rpr.append(rfonts)
    rfonts.set(qn('w:eastAsia'), CHINESE_FONT)
    rfonts.set(qn('w:ascii'), CHINESE_FONT)
    rfonts.set(qn('w:hAnsi'), CHINESE_FONT)
    # Page setup — A4 landscape to match the PDF orientation
    section = doc.sections[0]
    section.page_height, section.page_width = section.page_width, section.page_height
    section.left_margin = Inches(0.5)
    section.right_margin = Inches(0.5)
    section.top_margin = Inches(0.5)
    section.bottom_margin = Inches(0.5)
    return doc


# ---------------------------------------------------------------------------
# HTML -> DOCX conversion
# ---------------------------------------------------------------------------
# Severity colours (mirror Bootstrap classes used in templates)
SEVERITY_COLORS = {
    'danger': (0xC0, 0x39, 0x2B),   # red
    'warning': (0xE6, 0x9D, 0x00),  # amber
    'info': (0x29, 0x80, 0xB9),     # blue
    'success': (0x18, 0x9A, 0x4B),  # green
    'secure': (0x18, 0x9A, 0x4B),
    'high': (0xC0, 0x39, 0x2B),
    'medium': (0xE6, 0x9D, 0x00),
    'low': (0x29, 0x80, 0xB9),
    'hotspot': (0xFF, 0x4C, 0x4C),
}


def _class_color(node, default=None):
    """Look for severity / status colour classes on a node."""
    classes = node.get('class') or []
    if isinstance(classes, str):
        classes = classes.lower().split()
    else:
        classes = [str(c).lower() for c in classes]
    for cls in classes:
        if cls in SEVERITY_COLORS:
            return SEVERITY_COLORS[cls]
    return default


def _add_heading(doc, soup_node):
    text = soup_node.get_text(strip=True)
    if not text:
        return
    level_str = re.sub(r'[^0-9]', '', soup_node.name) or '2'
    try:
        level = max(1, min(6, int(level_str)))
    except ValueError:
        level = 2
    para = doc.add_heading(level=min(level, 4))
    run = para.add_run(text)
    _set_run_font(run, size=20 - 2 * (level - 1), bold=True)


def _add_paragraph(doc, soup_node):
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    color = _class_color(soup_node)
    # Detect alignment class
    classes = soup_node.get('class') or []
    if isinstance(classes, str):
        classes = classes.lower().split()
    else:
        classes = [str(c).lower() for c in classes]
    if 'text-center' in classes:
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if 'text-right' in classes:
        para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    _emit_inline(para, soup_node, color=color)


def _emit_inline(paragraph, soup_node, color=None):
    """Recursively walk children of a node and add runs / nested content."""
    for child in soup_node.children:
        if isinstance(child, NavigableString):
            text = unescape(str(child))
            if not text.strip() and not text:
                continue
            run = paragraph.add_run(text)
            _set_run_font(run, size=11, color=color)
            continue
        name = (child.name or '').lower()
        if name in ('br',):
            paragraph.add_run().add_break()
        elif name in ('strong', 'b'):
            text = child.get_text()
            run = paragraph.add_run(text)
            _set_run_font(run, size=11, bold=True, color=color)
        elif name in ('em', 'i'):
            text = child.get_text()
            run = paragraph.add_run(text)
            _set_run_font(run, size=11, color=color)
            run.italic = True
        elif name in ('code', 'span', 'small', 'mark'):
            run_color = _class_color(child, default=color)
            run = paragraph.add_run(child.get_text())
            _set_run_font(run, size=10, color=run_color)
        elif name in ('a',):
            text = child.get_text()
            run = paragraph.add_run(text)
            _set_run_font(run, size=11, color=color)
            run.underline = True
        elif name in ('img',):
            # Single image inside a paragraph - try to inline
            _try_add_image(paragraph, child)
        elif name in ('ul', 'ol'):
            _add_list(doc := paragraph.part.document, child)  # noqa: F821
        else:
            # Unknown inline - fallback to plain text
            run = paragraph.add_run(child.get_text())
            _set_run_font(run, size=11, color=color)


def _try_add_image(paragraph, img_tag):
    src = img_tag.get('src', '')
    img_bytes = _inline_image_bytes(src)
    if not img_bytes:
        return
    try:
        run = paragraph.add_run()
        run.add_picture(io.BytesIO(img_bytes), width=Inches(2.5))
    except Exception as exc:  # pragma: no cover - defensive
        logger.debug('Cannot inline image: %s', exc)


def _add_list(doc, soup_node):
    ordered = (soup_node.name or '').lower() == 'ol'
    for li in soup_node.find_all('li', recursive=False):
        para = doc.add_paragraph(style='List Number' if ordered else 'List Bullet')
        color = _class_color(li)
        _emit_inline(para, li, color=color)


def _add_table(doc, soup_node):
    rows = soup_node.find_all('tr')
    if not rows:
        return
    # Use first row as header if it has <th> in it
    first_cells = rows[0].find_all(['td', 'th'])
    has_header = any(c.name == 'th' for c in first_cells)
    data = []
    for r in rows:
        cells = r.find_all(['td', 'th'])
        data.append(cells)
    n_cols = max(len(r) for r in data) if data else 0
    if n_cols == 0:
        return
    table = doc.add_table(rows=len(data), cols=n_cols)
    table.style = 'Light Grid Accent 1'
    table.autofit = True
    for r_idx, row_cells in enumerate(data):
        for c_idx in range(n_cols):
            cell = table.cell(r_idx, c_idx)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
            if c_idx < len(row_cells):
                cell_html = row_cells[c_idx]
            else:
                cell_html = None
            # Reset default paragraph
            cell.paragraphs[0].text = ''
            if cell_html is None:
                continue
            is_header = (r_idx == 0 and has_header) or cell_html.name == 'th'
            color = _class_color(cell_html)
            _emit_inline(cell.paragraphs[0], cell_html, color=color)
            if is_header:
                for run in cell.paragraphs[0].runs:
                    _set_run_font(run, size=11, bold=True, color=color)


def _add_image_block(doc, img_tag):
    src = img_tag.get('src', '')
    img_bytes = _inline_image_bytes(src)
    if not img_bytes:
        return
    try:
        para = doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = para.add_run()
        run.add_picture(io.BytesIO(img_bytes), width=Inches(3.5))
    except Exception as exc:  # pragma: no cover
        logger.debug('Cannot inline image: %s', exc)


def _add_hr(doc):
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    from docx.oxml import OxmlElement
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), 'auto')
    pBdr.append(bottom)
    pPr.append(pBdr)


def _walk(doc, soup_node, depth=0):
    """Recursively walk an HTML node and emit DOCX content.

    The MobSF PDF templates use ``<article>`` and deeply nested
    ``<div class="container">`` blocks, so a simple top-level iteration
    over ``body.children`` is not enough: we have to descend into every
    container we encounter so that the underlying ``<h2>`` / ``<table>``
    / ``<p>`` nodes are reached.
    """
    if depth > 50:  # pragma: no cover - safety net
        return
    for node in list(soup_node.children):
        if isinstance(node, NavigableString):
            text = unescape(str(node)).strip()
            if not text:
                continue
            para = doc.add_paragraph()
            run = para.add_run(text)
            _set_run_font(run, size=11)
            continue
        name = (node.name or '').lower()
        if name in ('h1', 'h2', 'h3', 'h4', 'h5', 'h6'):
            _add_heading(doc, node)
        elif name == 'p':
            _add_paragraph(doc, node)
        elif name == 'img':
            _add_image_block(doc, node)
        elif name == 'table':
            _add_table(doc, node)
        elif name in ('ul', 'ol'):
            _add_list(doc, node)
        elif name == 'hr':
            _add_hr(doc)
        elif name in ('div', 'section', 'article', 'header', 'footer',
                      'main', 'aside', 'nav', 'body', 'html'):
            # Containers - just descend
            _walk(doc, node, depth + 1)
        elif name in ('style', 'script', 'head', 'meta', 'title', 'link',
                      'br'):
            # Skip
            continue
        else:
            # Unknown element - descend to be safe
            _walk(doc, node, depth + 1)


def html_to_docx(html_str, context=None):
    """Convert a rendered MobSF PDF report HTML string to a ``Document``."""
    if not DOCX_AVAILABLE:  # pragma: no cover
        raise RuntimeError(
            f'python-docx is not installed: {IMPORT_ERROR}')
    soup = BeautifulSoup(html_str, 'html.parser')

    doc = _bootstrap_document()
    body = soup.body or soup
    _walk(doc, body)
    return doc


# ---------------------------------------------------------------------------
# Django view
# ---------------------------------------------------------------------------
@login_required
def docx(request, checksum, api=False, jsonres=False):
    """Generate and stream a DOCX report for the given scan MD5."""
    if not DOCX_AVAILABLE:  # pragma: no cover
        return print_n_send_error_response(
            request,
            'python-docx is not installed on the server.',
            api,
        )
    try:
        if not is_md5(checksum):
            if api:
                return {'error': '无效的哈希'}
            return HttpResponse(
                json.dumps({'md5': '无效的哈希'}),
                content_type=ctype, status=500)

        android_static_db = StaticAnalyzerAndroid.objects.filter(MD5=checksum)
        ios_static_db = StaticAnalyzerIOS.objects.filter(MD5=checksum)
        win_static_db = StaticAnalyzerWindows.objects.filter(MD5=checksum)

        if android_static_db.exists():
            context, template = handle_pdf_android(android_static_db)
        elif ios_static_db.exists():
            context, template = handle_pdf_ios(ios_static_db)
        elif win_static_db.exists():
            context, template = handle_pdf_win(win_static_db)
        else:
            if api:
                return {'report': '未找到报告'}
            return HttpResponse(
                json.dumps({'report': '未找到报告'}),
                content_type=ctype, status=500)

        # VirusTotal (best effort)
        context['virus_total'] = None
        ext = os.path.splitext(context['file_name'].lower())[1]
        if settings.VT_ENABLED and ext != '.zip':
            app_bin = os.path.join(
                settings.UPLD_DIR,
                checksum + '/',
                checksum + ext)
            try:
                vt = VirusTotal.VirusTotal(checksum)
                context['virus_total'] = vt.get_result(app_bin)
            except Exception:
                logger.exception('VirusTotal lookup failed in DOCX generation')

        # Build the same file:// / local URL hints used by the PDF pipeline
        proto = 'file://'
        if platform.system() == 'Windows':
            proto = 'file:///'
        context['base_url'] = proto + settings.BASE_DIR
        context['dwd_dir'] = proto + settings.DWD_DIR
        context['host_os'] = 'nix' if platform.system() != 'Windows' else 'windows'
        context['timestamp'] = RecentScansDB.objects.get(
            MD5=checksum).TIMESTAMP

        if api and jsonres:
            return {'report_dat': context}

        html = template.render(context)
        document = html_to_docx(html, context=context)
        buf = io.BytesIO()
        document.save(buf)
        buf.seek(0)
        docx_bytes = buf.getvalue()
        if api:
            return {'docx_dat': docx_bytes}
        filename = '{}_{}_mobsf_report.docx'.format(
            context.get('file_name', 'report').rsplit('.', 1)[0],
            checksum,
        )
        response = HttpResponse(
            docx_bytes,
            content_type=(
                'application/vnd.openxmlformats-officedocument'
                '.wordprocessingml.document'),
        )
        response['Content-Disposition'] = (
            f'attachment; filename="{filename}"')
        return response
    except Exception as exp:
        logger.exception('Error Generating DOCX Report')
        msg = str(exp)
        if api:
            return print_n_send_error_response(request, msg, True, exp)
        return print_n_send_error_response(request, msg, False, exp)
